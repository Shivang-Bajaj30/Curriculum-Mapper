from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import json
import time
import tempfile

from skillToJob2 import skills_to_job_roles

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*", "methods": ["GET", "POST", "OPTIONS"], "allow_headers": ["*"]}})


def get_gemini_model():
    """Configure and return a Gemini model."""
    from dotenv import load_dotenv
    import google.generativeai as genai

    env_path = os.path.join(os.path.dirname(__file__), ".env")
    load_dotenv(env_path, override=True)
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY not set in backend/.env")
    genai.configure(api_key=api_key, transport="rest")
    return genai.GenerativeModel("gemini-2.5-flash")


def extract_text_from_file(filepath: str, filename: str) -> str:
    """Extract plain text from a PDF, DOCX, or TXT file."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        import fitz  # PyMuPDF
        text = ""
        doc = fitz.open(filepath)
        for page in doc:
            text += page.get_text()
        doc.close()
        return text

    elif ext == ".docx":
        from docx import Document
        doc = Document(filepath)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(lines)

    elif ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    return ""


def _strip_json_fence(response_text: str) -> str:
    response_text = response_text.strip()
    if response_text.startswith("```"):
        lines = response_text.split("\n")
        inner = []
        for line in lines[1:]:
            if line.strip() == "```":
                break
            inner.append(line)
        response_text = "\n".join(inner)
    return response_text.strip()


def extract_extracted_content_with_ai(text: str) -> dict:
    """
    Use Gemini to return structured JSON:
    { "extracted_content": [ { "topics": [ "skill1", "skill2", ... ] } ] }
    Topics = important technical / academic skills from the curriculum.
    """
    MAX_CHARS = 50000
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[Document truncated...]"

    model = get_gemini_model()

    prompt = f"""You are a curriculum and skills analyst. Read the curriculum document and extract the most important technical and transferable SKILLS and TOPICS
(students would list these on a resume or LinkedIn). Include tools, frameworks, languages, methods, and domain concepts.

Return ONLY valid JSON with this exact shape (no markdown, no explanation):
{{
  "extracted_content": [
    {{
      "topics": ["skill or topic 1", "skill or topic 2", "..."]
    }}
  ]
}}

Rules:
- Put ALL skills/topics in the single "topics" array inside the first object of "extracted_content".
- Aim for 15–40 distinct items when the document is rich; fewer if the document is short.
- Be specific (e.g. "Python", "REST APIs", "Machine Learning") not vague ("communication").

CURRICULUM TEXT:
---
{text}
---"""

    for attempt in range(3):
        try:
            response = model.generate_content(prompt)
            response_text = _strip_json_fence(response.text)
            data = json.loads(response_text)
            if not isinstance(data, dict):
                continue
            ec = data.get("extracted_content")
            if not isinstance(ec, list) or not ec:
                continue
            first = ec[0] if isinstance(ec[0], dict) else {}
            topics = first.get("topics")
            if not isinstance(topics, list):
                topics = []
            topics = [str(t).strip() for t in topics if t and str(t).strip()]
            data["extracted_content"] = [{"topics": topics}]
            return data
        except Exception as e:
            if attempt == 2:
                raise e
            time.sleep(2)

    return {"extracted_content": [{"topics": []}]}


def topics_flat_from_extracted(data: dict) -> list[str]:
    """Flatten topics from extracted_content for skill→job matching."""
    out = []
    ec = data.get("extracted_content") or []
    for block in ec:
        if isinstance(block, dict):
            t = block.get("topics") or []
            if isinstance(t, list):
                out.extend(str(x).strip() for x in t if x and str(x).strip())
    return out


@app.route("/")
def index():
    return "<p>Curriculum Mapper API</p>"


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "message": "Backend is running"})


@app.route("/api/upload", methods=["POST"])
def upload():
    """
    Upload curriculum files, extract text, then use AI to produce a clean topics list.
    Returns: { uploaded: [...], topics: [...], job_matches: [...], preview: "..." }
    """
    files = request.files.getlist("files") or request.files.getlist("file")
    uploaded_names = [f.filename for f in files if f.filename]

    if not uploaded_names:
        return jsonify({"error": "no files provided"}), 400

    all_text = ""
    for f in files:
        if not f.filename:
            continue
        suffix = os.path.splitext(f.filename)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            f.save(tmp.name)
            tmp_path = tmp.name
        try:
            text = extract_text_from_file(tmp_path, f.filename)
            all_text += f"\n\n=== {f.filename} ===\n{text}"
        finally:
            os.unlink(tmp_path)

    if not all_text.strip():
        return jsonify({
            "uploaded": uploaded_names,
            "extracted_content": [{"topics": []}],
            "topics": [],
            "job_matches": [],
            "preview": "",
            "warning": "Could not extract any text from the uploaded file(s).",
        })

    try:
        extracted = extract_extracted_content_with_ai(all_text)
    except Exception as e:
        return jsonify({
            "uploaded": uploaded_names,
            "extracted_content": [{"topics": []}],
            "topics": [],
            "job_matches": [],
            "preview": all_text[:300],
            "warning": f"AI extraction failed: {str(e)}",
        })

    topics = topics_flat_from_extracted(extracted)

    # --- Skill-to-Job matching via SentenceTransformer (skillToJob2.py) ---
    job_matches = []
    job_warning = None
    try:
        job_matches = skills_to_job_roles(topics, top_k=12)
    except Exception as je:
        job_warning = f"Job role matching unavailable: {je}"

    body = {
        "uploaded": uploaded_names,
        "extracted_content": extracted.get("extracted_content", [{"topics": topics}]),
        "topics": topics,
        "job_matches": job_matches,
        "preview": all_text[:500].strip(),
    }
    if job_warning:
        body["warning"] = job_warning
    return jsonify(body)


@app.route("/api/jobs", methods=["POST"])
def jobs():
    """
    Directly match a list of skills to job roles without needing a file upload.

    Request body (JSON):
        {
            "skills": ["Python", "Machine Learning", "SQL"],   // required
            "top_k": 10                                         // optional, default 12
        }

    Response:
        {
            "skills": [...],
            "job_matches": [
                {
                    "job_title": "Data Scientist",
                    "match_score": 0.8921,
                    "skills_required": "...",
                    "industry": "...",
                    "experience_level": "...",
                    "education_required": "..."
                },
                ...
            ]
        }
    """
    data = request.get_json(silent=True) or {}

    # Accept either "skills" (preferred) or "topics" (alias for compatibility)
    skills = data.get("skills") or data.get("topics") or []
    top_k = int(data.get("top_k", 12))

    if not skills:
        return jsonify({"error": "skills list is required and must be non-empty"}), 400

    if not isinstance(skills, list):
        return jsonify({"error": "skills must be a JSON array of strings"}), 400

    # Sanitise: ensure all entries are non-empty strings
    skills = [str(s).strip() for s in skills if s and str(s).strip()]
    if not skills:
        return jsonify({"error": "No valid skill strings found after sanitising input"}), 400

    try:
        job_matches = skills_to_job_roles(skills, top_k=top_k)
    except FileNotFoundError as e:
        return jsonify({"error": str(e)}), 500
    except Exception as e:
        return jsonify({"error": f"Job matching failed: {str(e)}"}), 500

    return jsonify({
        "skills": skills,
        "job_matches": job_matches,
    })


@app.route("/api/analyze", methods=["POST"])
def analyze():
    """
    Body: { query: string, topics: string[] }
    Returns a compatibility analysis: summary paragraph + per-topic pointers.
    """
    data = request.get_json(silent=True) or {}
    query = (data.get("query") or "").strip()
    topics = data.get("topics") or []

    if not query:
        return jsonify({"error": "query is required"}), 400

    if not topics:
        return jsonify({"error": "No curriculum topics found. Please upload a curriculum file first."}), 400

    try:
        model = get_gemini_model()

        topics_text = "\n".join(f"- {t}" for t in topics)

        prompt = f"""You are a curriculum-to-career compatibility expert.

A student is exploring whether this course matches their target role or skills.

Target Role / Keywords: "{query}"

Course Topics Extracted from Curriculum:
{topics_text}

Produce a structured compatibility analysis. Return ONLY valid JSON with this exact structure:
{{
  "summary": "A concise 2-3 sentence paragraph describing how well the course aligns with the target role and what the learner will gain.",
  "compatibility_score": <integer 0-100 representing overall compatibility>,
  "pointers": [
    {{
      "topic": "exact course topic from the list above",
      "relevance": "brief 1-2 sentence explanation of how this topic directly applies to the role",
      "match_level": "high|medium|low"
    }}
  ]
}}

Rules:
- Include only the most relevant topics (top 8 maximum)
- Order by relevance descending (high first)
- Be specific and practical in the relevance explanations
- Return ONLY valid JSON, no markdown, no explanation"""

        response = model.generate_content(prompt)
        response_text = response.text.strip()

        if response_text.startswith("```"):
            lines = response_text.split("\n")
            inner = []
            for line in lines[1:]:
                if line.strip() == "```":
                    break
                inner.append(line)
            response_text = "\n".join(inner)

        result = json.loads(response_text)
        return jsonify({"success": True, "analysis": result})

    except json.JSONDecodeError as e:
        return jsonify({"error": f"AI response could not be parsed: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/search", methods=["POST"])
def search():
    data = request.get_json(silent=True) or {}
    query = (data.get("query") or "").strip()
    if not query:
        return jsonify({"error": "query is required"}), 400
    return jsonify({"query": query, "results": []})